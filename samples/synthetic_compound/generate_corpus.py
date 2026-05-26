"""Generator for the synthetic IB source corpus.

Produces a realistic-shaped set of source documents (DOCX study reports +
XLSX data tables) for a fictional kinase-inhibitor compound. Output is
written under `sources/` and is deterministic — re-running produces the
same bytes (no random data).

The corpus is deliberately compact (~15 documents, each <2k words) so it
fits comfortably in an LLM context window during PoC generation runs.
All data is synthetic; no real compound, patient, or study.

Run:
    python samples/synthetic_compound/generate_corpus.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook


HERE = Path(__file__).parent
OUT = HERE / "sources"
METADATA = json.loads((HERE / "metadata.json").read_text(encoding="utf-8"))
COMPOUND = METADATA["compound"]["research_code"]


# -- Document helpers --------------------------------------------------------


def _new_doc(title: str, study_id: str | None = None) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    doc.add_heading(title, level=1)
    if study_id:
        p = doc.add_paragraph()
        p.add_run(f"Study ID: {study_id}").bold = True
        p.add_run(f"\nCompound: {COMPOUND} ({METADATA['compound']['sponsor_code']})")
    return doc


def _save(doc: Document, relpath: str) -> Path:
    target = OUT / relpath
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target


def _table_2col(doc: Document, header: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = header[0]
    table.rows[0].cells[1].text = header[1]
    for i, (a, b) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b


# -- Documents ---------------------------------------------------------------


def cmc_summary() -> Path:
    doc = _new_doc("CMC Summary — XYZ-001")
    doc.add_heading("Chemical Identity", level=2)
    doc.add_paragraph(
        f"{COMPOUND} is a small-molecule selective inhibitor of Kinase Z, "
        "developed as an oral therapy for solid tumors expressing Kinase Z. "
        "Chemical name (synthetic): 4-[(3-chloro-4-fluorophenyl)amino]-7-methoxy-6-"
        "(piperidin-4-yloxy)quinazoline. Molecular formula: C20H21ClFN5O3. "
        "Molecular weight: 433.86 g/mol."
    )

    doc.add_heading("Physicochemical Properties", level=2)
    _table_2col(
        doc,
        ("Property", "Value"),
        [
            ("Appearance", "White to off-white crystalline powder"),
            ("Solubility (water, 25 °C)", "0.12 mg/mL"),
            ("Solubility (DMSO)", ">100 mg/mL"),
            ("Log P", "3.4"),
            ("pKa", "8.1"),
            ("Polymorphic forms", "Form I (anhydrous) used for clinical material"),
        ],
    )

    doc.add_heading("Formulation", level=2)
    doc.add_paragraph(
        "Clinical formulation is a film-coated immediate-release tablet containing "
        "25 mg or 100 mg of XYZ-001 as the free base. Excipients include "
        "microcrystalline cellulose, lactose monohydrate, croscarmellose sodium, "
        "magnesium stearate, and a standard Opadry coating. Tablets are stored at "
        "15-25 °C in HDPE bottles with desiccant."
    )

    doc.add_heading("Stability", level=2)
    doc.add_paragraph(
        "Long-term stability data through 24 months at 25 °C / 60% RH supports the "
        "current shelf life of 24 months. No significant degradation or impurity "
        "growth has been observed."
    )
    return _save(doc, "cmc/cmc_summary.docx")


def primary_pharmacology() -> Path:
    doc = _new_doc("Primary Pharmacology of XYZ-001", study_id="XYZ-NC-001")
    doc.add_heading("Mechanism of Action", level=2)
    doc.add_paragraph(
        "XYZ-001 is a reversible ATP-competitive inhibitor of Kinase Z. In a panel of "
        "468 human kinases, XYZ-001 inhibited Kinase Z with an IC50 of 3.2 nM, with "
        ">100-fold selectivity over the closest off-target kinase (Kinase Y, IC50 410 nM)."
    )

    doc.add_heading("In Vitro Efficacy", level=2)
    doc.add_paragraph(
        "XYZ-001 inhibited proliferation of Kinase Z-amplified tumor cell lines "
        "(HCC-827, NCI-H1975 surrogates) with mean IC50 of 18 nM. In wild-type lines "
        "lacking Kinase Z amplification, IC50 was >5,000 nM, supporting target-driven "
        "antiproliferative activity."
    )

    doc.add_heading("In Vivo Efficacy", level=2)
    doc.add_paragraph(
        "In a mouse xenograft model of Kinase Z-amplified tumor (mean tumor volume "
        "200 mm³ at randomization), oral XYZ-001 at 10, 30, and 100 mg/kg QD for 21 days "
        "produced tumor growth inhibition (TGI) of 42%, 78%, and 96% respectively versus "
        "vehicle (n=10 per group, p<0.001 for 30 and 100 mg/kg doses by two-way ANOVA). "
        "Complete regressions were observed in 4/10 animals at 100 mg/kg."
    )

    doc.add_heading("Safety Pharmacology", level=2)
    doc.add_paragraph(
        "In a hERG assay, XYZ-001 inhibited the hERG potassium channel with an IC50 of "
        "12 µM, providing approximately 60-fold margin over the projected human Cmax at "
        "the anticipated efficacious dose. In an anesthetized dog cardiovascular study "
        "(single doses up to 30 mg/kg IV), no clinically meaningful changes in heart rate, "
        "blood pressure, or QTc interval were observed."
    )
    return _save(doc, "nonclinical/pharmacology/XYZ-NC-001_primary_pharmacology.docx")


def nonclinical_pk_rat() -> Path:
    doc = _new_doc("Nonclinical Pharmacokinetics in Rat", study_id="XYZ-NC-PK-001")
    doc.add_heading("Study Design", level=2)
    doc.add_paragraph(
        "Sprague-Dawley rats (n=4/sex/dose) received XYZ-001 as a single oral dose "
        "of 5, 25, or 100 mg/kg, or by intravenous bolus at 5 mg/kg. Plasma samples "
        "were collected through 48 hours post-dose and analyzed by LC-MS/MS."
    )

    doc.add_heading("Pharmacokinetic Parameters — Oral Dosing", level=2)
    _table_2col(
        doc,
        ("Parameter", "Value (mean, 100 mg/kg PO)"),
        [
            ("Cmax", "1,840 ng/mL"),
            ("Tmax", "2.0 h"),
            ("AUC0-inf", "11,400 ng·h/mL"),
            ("Half-life (t1/2)", "8.6 h"),
            ("Apparent oral clearance (CL/F)", "8.8 L/h/kg"),
            ("Apparent volume of distribution (Vz/F)", "109 L/kg"),
        ],
    )

    doc.add_heading("Dose Linearity", level=2)
    doc.add_paragraph(
        "Cmax and AUC increased approximately dose-proportionally across the 5–100 mg/kg "
        "range with no evidence of saturation. Absolute oral bioavailability calculated "
        "from the 5 mg/kg IV comparator was 38%."
    )

    doc.add_heading("Tissue Distribution", level=2)
    doc.add_paragraph(
        "In a separate quantitative whole-body autoradiography study, [14C]-XYZ-001 "
        "distributed broadly with highest concentrations in liver, kidney, and tumor "
        "tissue at 4 hours post-dose. Brain penetration was minimal (brain:plasma ratio "
        "0.04)."
    )
    return _save(doc, "nonclinical/pk/XYZ-NC-PK-001_pk_rat.docx")


def nonclinical_pk_dog() -> Path:
    doc = _new_doc("Nonclinical Pharmacokinetics in Dog", study_id="XYZ-NC-PK-002")
    doc.add_heading("Study Design", level=2)
    doc.add_paragraph(
        "Beagle dogs (n=3/sex/dose) received XYZ-001 as a single oral dose of 1, 5, "
        "or 25 mg/kg, or by intravenous bolus at 1 mg/kg. Plasma samples were "
        "collected through 72 hours post-dose."
    )

    doc.add_heading("Pharmacokinetic Parameters — Oral Dosing", level=2)
    _table_2col(
        doc,
        ("Parameter", "Value (mean, 25 mg/kg PO)"),
        [
            ("Cmax", "920 ng/mL"),
            ("Tmax", "3.0 h"),
            ("AUC0-inf", "8,600 ng·h/mL"),
            ("Half-life (t1/2)", "14.2 h"),
            ("Apparent oral clearance (CL/F)", "2.9 L/h/kg"),
            ("Apparent volume of distribution (Vz/F)", "59 L/kg"),
            ("Absolute oral bioavailability", "62%"),
        ],
    )
    doc.add_heading("Dose Linearity", level=2)
    doc.add_paragraph(
        "Dose-proportional exposure was observed across 1–25 mg/kg in dog. Half-life was "
        "longer in dog than rat, consistent with lower body-weight-normalized clearance."
    )
    return _save(doc, "nonclinical/pk/XYZ-NC-PK-002_pk_dog.docx")


def nonclinical_tox_rat_13wk() -> Path:
    doc = _new_doc("13-Week Repeat-Dose Toxicology in Rat", study_id="XYZ-NC-004")
    doc.add_heading("Study Design", level=2)
    doc.add_paragraph(
        "Sprague-Dawley rats (n=15/sex/group, plus 5/sex/group recovery) received oral "
        "doses of vehicle, 10, 30, or 100 mg/kg/day for 13 weeks, followed by a 4-week "
        "recovery period for the high-dose and control groups."
    )

    doc.add_heading("Mortality and Clinical Signs", level=2)
    doc.add_paragraph(
        "No mortality occurred in vehicle, 10, or 30 mg/kg/day groups. At 100 mg/kg/day, "
        "2 of 30 animals were euthanized in moribund condition on study days 47 and 62 "
        "due to severe weight loss and decreased activity. Surviving high-dose animals "
        "exhibited reduced food consumption and body weight gain (mean −12% versus vehicle "
        "at week 13)."
    )

    doc.add_heading("Clinical Pathology", level=2)
    doc.add_paragraph(
        "At 100 mg/kg/day, dose-related elevations in ALT (mean 2.4-fold) and AST (mean "
        "1.9-fold) were observed at week 13. Mild decreases in red cell mass parameters "
        "(hemoglobin −8%, hematocrit −7%) were noted at 100 mg/kg/day. All clinical "
        "pathology changes reversed during the 4-week recovery period."
    )

    doc.add_heading("Anatomic Pathology", level=2)
    doc.add_paragraph(
        "Microscopic findings at 100 mg/kg/day included minimal-to-mild hepatocellular "
        "hypertrophy (10/15 males, 7/15 females) and minimal bone marrow erythroid "
        "hypocellularity (6/15 males). These findings reversed during recovery. No "
        "microscopic findings were attributed to XYZ-001 at 10 or 30 mg/kg/day."
    )

    doc.add_heading("NOAEL Determination", level=2)
    p = doc.add_paragraph()
    p.add_run("NOAEL: 30 mg/kg/day").bold = True
    p.add_run(
        ". The NOAEL was based on the absence of clinical-pathology and microscopic "
        "findings at this dose, with minimal hepatocellular hypertrophy and erythroid "
        "effects at 100 mg/kg/day. Mean AUC0-24 at the NOAEL was 6,200 ng·h/mL in males "
        "and 7,100 ng·h/mL in females."
    )
    return _save(doc, "nonclinical/toxicology/XYZ-NC-004_tox_rat_13wk.docx")


def nonclinical_tox_dog_26wk() -> Path:
    doc = _new_doc("26-Week Repeat-Dose Toxicology in Dog", study_id="XYZ-NC-005")
    doc.add_heading("Study Design", level=2)
    doc.add_paragraph(
        "Beagle dogs (n=4/sex/group) received oral doses of vehicle, 2, 6, or 20 mg/kg/day "
        "for 26 weeks. A 4-week recovery cohort (n=2/sex) was included for the high-dose "
        "and control groups."
    )

    doc.add_heading("Mortality and Clinical Signs", level=2)
    doc.add_paragraph(
        "No unscheduled deaths occurred. Animals at 20 mg/kg/day exhibited intermittent "
        "soft stools (4/8 animals) starting in week 6. No effects on body weight or food "
        "consumption were observed at any dose level."
    )

    doc.add_heading("Clinical Pathology", level=2)
    doc.add_paragraph(
        "At 20 mg/kg/day, mild and reversible elevations in ALT (mean 1.7-fold) were "
        "observed at week 13 and resolved by week 26. No other clinically meaningful "
        "clinical-pathology changes were noted at any dose."
    )

    doc.add_heading("Cardiovascular Assessment", level=2)
    doc.add_paragraph(
        "ECGs collected at weeks 4, 13, and 26 showed no XYZ-001-related changes in heart "
        "rate, PR, QRS, or QTc intervals at any dose level. QTcF interval at 20 mg/kg/day "
        "was within 1.5% of baseline."
    )

    doc.add_heading("NOAEL Determination", level=2)
    p = doc.add_paragraph()
    p.add_run("NOAEL: 6 mg/kg/day").bold = True
    p.add_run(
        ". The NOAEL was based on the soft-stool clinical signs and mild reversible ALT "
        "elevation at 20 mg/kg/day. Mean AUC0-24 at the NOAEL was 4,800 ng·h/mL."
    )
    return _save(doc, "nonclinical/toxicology/XYZ-NC-005_tox_dog_26wk.docx")


def clinical_pk_study() -> Path:
    doc = _new_doc("Clinical Pharmacokinetics — Study XYZ-101", study_id="XYZ-101")
    doc.add_heading("Study Design", level=2)
    doc.add_paragraph(
        "XYZ-101 was a Phase 1 dose-escalation study of oral XYZ-001 in 36 patients with "
        "advanced solid tumors expressing Kinase Z. Six dose cohorts (25, 50, 100, 200, "
        "300, and 400 mg QD) were evaluated using a 3+3 design with intensive PK sampling "
        "on Day 1 and Day 28."
    )

    doc.add_heading("Single-Dose Pharmacokinetics", level=2)
    _table_2col(
        doc,
        ("Parameter (mean, 200 mg QD)", "Day 1 Value"),
        [
            ("Cmax", "612 ng/mL"),
            ("Tmax", "2.5 h"),
            ("AUC0-24", "5,400 ng·h/mL"),
            ("AUC0-inf", "7,800 ng·h/mL"),
            ("Apparent t1/2", "16.2 h"),
            ("Apparent CL/F", "25.6 L/h"),
        ],
    )

    doc.add_heading("Steady-State Pharmacokinetics", level=2)
    doc.add_paragraph(
        "Steady-state was reached by Day 7 of QD dosing. The accumulation ratio (Day 28 "
        "AUC0-24 / Day 1 AUC0-24) was 1.6, consistent with the observed half-life. "
        "Cmax and AUC at steady state increased approximately dose-proportionally from "
        "25 to 300 mg, with slight overproportionality observed at 400 mg."
    )

    doc.add_heading("Effect of Food", level=2)
    doc.add_paragraph(
        "A subset (n=12) crossed over to evaluate the effect of a high-fat meal on a "
        "200 mg dose. The high-fat meal increased Cmax by approximately 35% and AUC by "
        "approximately 22%. Patients were instructed to take XYZ-001 at least 2 hours "
        "before or after meals."
    )
    return _save(doc, "clinical/pk/XYZ-101_clinical_pk.docx")


def clinical_safety_study() -> Path:
    doc = _new_doc("Clinical Safety Summary — Studies XYZ-101 and XYZ-102", study_id="XYZ-101 / XYZ-102")
    doc.add_heading("Patient Exposure", level=2)
    doc.add_paragraph(
        "As of the data cutoff (March 2026), 120 patients had received at least one dose "
        "of XYZ-001 across studies XYZ-101 (n=36) and XYZ-102 (n=84). Median duration of "
        "treatment was 18.4 weeks (range 0.5–62.0 weeks). Sixty-two patients (52%) "
        "received treatment for ≥12 weeks."
    )

    doc.add_heading("Adverse Events Overview", level=2)
    doc.add_paragraph(
        "All 120 patients (100%) experienced at least one treatment-emergent adverse "
        "event (TEAE). The most frequent TEAEs (≥20% incidence) were diarrhea (58%), "
        "fatigue (44%), rash (38%), nausea (29%), and AST/ALT elevation (24%). Grade 3-4 "
        "TEAEs occurred in 41 patients (34%); the most common were ALT increased (8%), "
        "diarrhea (6%), and rash (5%)."
    )

    doc.add_heading("Serious Adverse Events", level=2)
    doc.add_paragraph(
        "Twenty-four patients (20%) experienced at least one serious adverse event (SAE). "
        "The most common SAEs were pneumonitis (n=4, 3%), pulmonary embolism (n=3, 2.5%), "
        "and dehydration (n=3, 2.5%). Four cases of pneumonitis included one grade 3 "
        "event leading to permanent discontinuation; the remaining three resolved with "
        "corticosteroid therapy."
    )

    doc.add_heading("Dose Modifications and Discontinuations", level=2)
    doc.add_paragraph(
        "Dose interruptions occurred in 47 patients (39%), most commonly for diarrhea "
        "(n=14), rash (n=11), or ALT elevation (n=9). Twelve patients (10%) permanently "
        "discontinued XYZ-001 due to a TEAE, including pneumonitis (n=2), ALT increased "
        "(n=2), rash (n=2), and individual events of pulmonary embolism, fatigue, and "
        "cardiac failure (each n=1)."
    )

    doc.add_heading("Deaths", level=2)
    doc.add_paragraph(
        "Eight patients died on study or within 30 days of the last XYZ-001 dose. Seven "
        "deaths were attributed to disease progression. One death (sudden cardiac death "
        "in a patient with prior cardiac history) was classified as possibly related to "
        "XYZ-001 by the investigator."
    )

    doc.add_heading("Preliminary Efficacy", level=2)
    doc.add_paragraph(
        "Among 78 patients evaluable for response in XYZ-102, the objective response rate "
        "(ORR) was 31% (24/78, 95% CI: 21-43). Median duration of response was 7.2 months "
        "(95% CI: 5.8-not reached). Median progression-free survival was 5.1 months "
        "(95% CI: 4.2-6.8)."
    )
    return _save(doc, "clinical/safety/XYZ-101_XYZ-102_safety_summary.docx")


def background_target_rationale() -> Path:
    doc = _new_doc("Background and Target Rationale — Kinase Z")
    doc.add_heading("Kinase Z Biology", level=2)
    doc.add_paragraph(
        "Kinase Z is a receptor tyrosine kinase implicated in the proliferation and "
        "survival of approximately 12% of solid tumors. Gene amplification of KINZ "
        "and activating point mutations (most commonly p.L883R) drive constitutive "
        "kinase activity and downstream MAPK and PI3K-AKT signaling."
    )

    doc.add_heading("Therapeutic Hypothesis", level=2)
    doc.add_paragraph(
        "Selective inhibition of Kinase Z is hypothesized to provide tumor regression "
        "in Kinase Z-altered cancers without the off-target toxicity associated with "
        "broad-spectrum kinase inhibitors. The therapeutic window depends on achieving "
        "sustained target coverage at clinically tolerable exposure."
    )

    doc.add_heading("Competitive Landscape", level=2)
    doc.add_paragraph(
        "Two other Kinase Z inhibitors are in clinical development (Compound A in "
        "Phase 1, Compound B in Phase 2). XYZ-001 is distinguished by higher kinase "
        "selectivity (>100-fold) and longer human half-life (16 h versus 6 h for "
        "Compound A)."
    )
    return _save(doc, "nonclinical/pharmacology/background_kinase_z.docx")


# -- Data tables (XLSX) -----------------------------------------------------


def ae_summary_xlsx() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "AE_by_SOC"
    ws.append(["MedDRA SOC", "Any Grade n (%)", "Grade 3-4 n (%)", "SAE n (%)"])
    rows = [
        ["Gastrointestinal disorders", "78 (65%)", "12 (10%)", "4 (3%)"],
        ["General disorders", "63 (53%)", "5 (4%)", "2 (2%)"],
        ["Skin and subcutaneous tissue disorders", "52 (43%)", "7 (6%)", "1 (1%)"],
        ["Investigations (lab abnormalities)", "47 (39%)", "13 (11%)", "2 (2%)"],
        ["Respiratory, thoracic and mediastinal disorders", "29 (24%)", "5 (4%)", "5 (4%)"],
        ["Vascular disorders", "18 (15%)", "4 (3%)", "3 (3%)"],
        ["Cardiac disorders", "9 (8%)", "2 (2%)", "2 (2%)"],
        ["Nervous system disorders", "31 (26%)", "3 (3%)", "1 (1%)"],
    ]
    for r in rows:
        ws.append(r)

    ws2 = wb.create_sheet("Top_AEs_PT")
    ws2.append(["Preferred Term", "Any Grade n (%)", "Grade 3-4 n (%)"])
    pt_rows = [
        ["Diarrhoea", "69 (58%)", "7 (6%)"],
        ["Fatigue", "53 (44%)", "3 (3%)"],
        ["Rash", "45 (38%)", "6 (5%)"],
        ["Nausea", "35 (29%)", "1 (1%)"],
        ["Alanine aminotransferase increased", "29 (24%)", "10 (8%)"],
        ["Aspartate aminotransferase increased", "26 (22%)", "4 (3%)"],
        ["Decreased appetite", "22 (18%)", "1 (1%)"],
        ["Vomiting", "18 (15%)", "2 (2%)"],
        ["Pneumonitis", "6 (5%)", "3 (3%)"],
    ]
    for r in pt_rows:
        ws2.append(r)

    target = OUT / "clinical" / "data" / "ae_summary.xlsx"
    target.parent.mkdir(parents=True, exist_ok=True)
    wb.save(target)
    return target


def exposure_summary_xlsx() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Exposure"
    ws.append(["Study", "Dose Group", "N", "Median Duration (weeks)", "Patients ≥12 wk"])
    rows = [
        ["XYZ-101", "25 mg QD", 3, "8.0", 1],
        ["XYZ-101", "50 mg QD", 3, "10.0", 2],
        ["XYZ-101", "100 mg QD", 6, "14.0", 4],
        ["XYZ-101", "200 mg QD", 6, "20.0", 5],
        ["XYZ-101", "300 mg QD", 12, "22.0", 8],
        ["XYZ-101", "400 mg QD", 6, "16.0", 4],
        ["XYZ-102", "300 mg QD (RP2D)", 84, "18.5", 38],
    ]
    for r in rows:
        ws.append(r)

    target = OUT / "clinical" / "data" / "exposure_summary.xlsx"
    target.parent.mkdir(parents=True, exist_ok=True)
    wb.save(target)
    return target


def pk_summary_xlsx() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Nonclinical_PK"
    ws.append(["Species", "Dose (mg/kg)", "Route", "Cmax (ng/mL)", "AUC0-inf (ng·h/mL)", "t1/2 (h)", "F (%)"])
    rows = [
        ["Rat", 5, "PO", 95, 580, 7.2, 38],
        ["Rat", 25, "PO", 480, 2900, 8.1, 38],
        ["Rat", 100, "PO", 1840, 11400, 8.6, 38],
        ["Dog", 1, "PO", 38, 340, 13.5, 62],
        ["Dog", 5, "PO", 185, 1700, 14.0, 62],
        ["Dog", 25, "PO", 920, 8600, 14.2, 62],
    ]
    for r in rows:
        ws.append(r)

    ws2 = wb.create_sheet("Clinical_PK_SS")
    ws2.append(["Study", "Dose", "N", "Cmax SS (ng/mL)", "AUC0-24 SS (ng·h/mL)", "Acc ratio"])
    crows = [
        ["XYZ-101", "25 mg QD", 3, 72, 590, 1.4],
        ["XYZ-101", "100 mg QD", 6, 304, 2680, 1.5],
        ["XYZ-101", "200 mg QD", 6, 612, 5400, 1.6],
        ["XYZ-101", "300 mg QD", 12, 945, 8400, 1.6],
        ["XYZ-101", "400 mg QD", 6, 1380, 12100, 1.7],
    ]
    for r in crows:
        ws2.append(r)

    target = OUT / "data" / "pk_summary.xlsx"
    target.parent.mkdir(parents=True, exist_ok=True)
    wb.save(target)
    return target


def pivotal_tox_summary_xlsx() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Pivotal_Tox"
    ws.append([
        "Study ID", "Species", "Duration", "Doses (mg/kg/day)",
        "NOAEL (mg/kg/day)", "NOAEL AUC0-24 (ng·h/mL)", "Target Organs of Toxicity",
    ])
    rows = [
        ["XYZ-NC-002", "Rat", "28-day", "0/20/60/200", "60", "12000", "Liver (hypertrophy)"],
        ["XYZ-NC-003", "Dog", "28-day", "0/4/12/40", "12", "9600", "GI tract (soft stools)"],
        ["XYZ-NC-004", "Rat", "13-week", "0/10/30/100", "30", "6200",
         "Liver (hypertrophy), bone marrow (erythroid)"],
        ["XYZ-NC-005", "Dog", "26-week", "0/2/6/20", "6", "4800",
         "GI tract (soft stools), liver (mild ALT elevation)"],
    ]
    for r in rows:
        ws.append(r)

    target = OUT / "data" / "pivotal_tox_summary.xlsx"
    target.parent.mkdir(parents=True, exist_ok=True)
    wb.save(target)
    return target


# -- Entry point ------------------------------------------------------------


def main() -> None:
    OUT.mkdir(exist_ok=True)
    paths: list[Path] = [
        cmc_summary(),
        background_target_rationale(),
        primary_pharmacology(),
        nonclinical_pk_rat(),
        nonclinical_pk_dog(),
        nonclinical_tox_rat_13wk(),
        nonclinical_tox_dog_26wk(),
        clinical_pk_study(),
        clinical_safety_study(),
        ae_summary_xlsx(),
        exposure_summary_xlsx(),
        pk_summary_xlsx(),
        pivotal_tox_summary_xlsx(),
    ]
    print(f"Generated {len(paths)} source documents under {OUT}")
    for p in paths:
        print(f"  {p.relative_to(OUT.parent.parent)}")


if __name__ == "__main__":
    main()
