"""Tests for the SampleReportsAdapter."""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from services.template_service import (
    SampleReportsAdapter,
    SampleReportsAdapterOptions,
    TemplateBuilder,
)
from shared.schemas import GenerationMode, TemplateStatus


def _make_sample(headings_with_content: list[tuple[int, str, str]]) -> bytes:
    """Build an in-memory .docx with given (level, heading, body) triples."""
    doc = Document()
    for level, heading, body in headings_with_content:
        doc.add_heading(heading, level=level)
        if body:
            doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Three "completed reports" that share most sections + have a few divergent ones.
SAMPLE_A = [
    (1, "1. Introduction", "Compound A was assessed in trial XA-101."),
    (1, "2. Nonclinical Studies", ""),
    (2, "2.1 Pharmacology", "Mechanism studies in vitro."),
    (2, "2.2 Toxicology", "13-week rat tox, NOAEL 30 mg/kg/day."),
    (1, "3. Effects in Humans", "Phase 1 dose escalation."),
    (1, "4. Summary", "Conclusions."),
]
SAMPLE_B = [
    (1, "Introduction", "Compound B was investigated in study XB-201."),  # no number prefix
    (1, "Nonclinical Studies", ""),
    (2, "Pharmacology", "Receptor binding profile."),
    (2, "Toxicology", "26-week dog tox, NOAEL 6 mg/kg/day."),
    (1, "Effects in Humans", "Phase 2 expansion."),
    (1, "References", "Bibliography only."),  # divergent: B has refs, others don't
]
SAMPLE_C = [
    (1, "1. Introduction", "Compound C, a kinase inhibitor."),
    (1, "2. Nonclinical Studies", ""),
    (2, "2.1 Pharmacology", "Off-target screening."),
    (2, "2.2 Toxicology", "Pivotal tox per ICH M3."),
    (2, "2.3 Pharmacokinetics", "Rat/dog PK."),  # divergent: only C has this
    (1, "3. Effects in Humans", "Phase 1/2 results."),
    (1, "4. Summary", "Risk/benefit."),
]


def test_adapter_extracts_common_sections(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    a.write_bytes(_make_sample(SAMPLE_A))
    b = tmp_path / "b.docx"
    b.write_bytes(_make_sample(SAMPLE_B))
    c = tmp_path / "c.docx"
    c.write_bytes(_make_sample(SAMPLE_C))

    adapter = SampleReportsAdapter(
        SampleReportsAdapterOptions(
            template_id="derived_ib",
            title="Derived IB template",
            min_occurrence_ratio=0.5,
        )
    )
    template = adapter.from_files([a, b, c])

    titles = [s.title.lower() for s in template.all_sections()]
    # Appears in all three (or 2/3 after normalization across "1. Introduction"
    # vs "Introduction") — must be kept.
    assert any("introduction" in t for t in titles)
    assert any("nonclinical studies" in t for t in titles)
    assert any("pharmacology" in t for t in titles)
    assert any("toxicology" in t for t in titles)
    assert any("effects in humans" in t for t in titles)

    # "References" only appears in sample B (1/3) — below the 0.5 threshold,
    # should be excluded.
    assert not any("references" in t for t in titles)
    # "Pharmacokinetics" only in sample C — also excluded.
    assert not any("pharmacokinetics" in t for t in titles)


def test_adapter_preserves_hierarchy(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    a.write_bytes(_make_sample(SAMPLE_A))
    c = tmp_path / "c.docx"
    c.write_bytes(_make_sample(SAMPLE_C))

    adapter = SampleReportsAdapter(
        SampleReportsAdapterOptions(template_id="t", min_occurrence_ratio=0.5)
    )
    template = adapter.from_files([a, c])
    # Pharmacology and Toxicology appear in both samples as children of
    # "Nonclinical Studies" — should be nested.
    parent = next(
        s for s in template.sections if "nonclinical" in s.title.lower()
    )
    children_titles = {c.title.lower() for c in parent.children}
    assert "pharmacology" in children_titles
    assert "toxicology" in children_titles


def test_adapter_emits_draft_status(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    a.write_bytes(_make_sample(SAMPLE_A))
    template = SampleReportsAdapter(
        SampleReportsAdapterOptions(template_id="t")
    ).from_files([a])
    assert template.status == TemplateStatus.DRAFT
    assert template.metadata.source_origin == "from_samples"


def test_adapter_uses_sample_content_in_prompt(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    a.write_bytes(_make_sample(SAMPLE_A))
    b = tmp_path / "b.docx"
    b.write_bytes(_make_sample(SAMPLE_B))

    template = SampleReportsAdapter(
        SampleReportsAdapterOptions(
            template_id="t",
            min_occurrence_ratio=0.5,
            max_hint_chars_per_section=200,
        )
    ).from_files([a, b])
    intro = next(s for s in template.all_sections() if "introduction" in s.title.lower())
    # The sample text should surface in the prompt as a content hint
    assert "Compound" in (intro.generation.prompt_template or "")


def test_adapter_handles_threshold_one(tmp_path: Path) -> None:
    """min_occurrence_ratio=1.0 requires every sample to have the section."""
    a = tmp_path / "a.docx"
    a.write_bytes(_make_sample(SAMPLE_A))
    b = tmp_path / "b.docx"
    b.write_bytes(_make_sample(SAMPLE_B))
    c = tmp_path / "c.docx"
    c.write_bytes(_make_sample(SAMPLE_C))

    template = SampleReportsAdapter(
        SampleReportsAdapterOptions(template_id="t", min_occurrence_ratio=1.0)
    ).from_files([a, b, c])
    titles = [s.title.lower() for s in template.all_sections()]
    # "Summary" appears in A and C but not B — excluded under strict mode.
    assert not any(t == "summary" for t in titles)
    # Cross-sample core sections still included.
    assert any("introduction" in t for t in titles)


def test_adapter_all_llm_sections_require_citations(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    a.write_bytes(_make_sample(SAMPLE_A))
    template = SampleReportsAdapter(
        SampleReportsAdapterOptions(template_id="t")
    ).from_files([a])
    for section in template.all_sections():
        if section.generation.mode in (GenerationMode.LLM, GenerationMode.HYBRID):
            assert section.citation_policy.required


def test_adapter_raises_on_empty_input() -> None:
    with pytest.raises(ValueError, match="at least one sample"):
        SampleReportsAdapter(
            SampleReportsAdapterOptions(template_id="t")
        ).from_files([])


# --- Builder façade --------------------------------------------------------


def test_builder_from_samples_directory(tmp_path: Path) -> None:
    (tmp_path / "a.docx").write_bytes(_make_sample(SAMPLE_A))
    (tmp_path / "b.docx").write_bytes(_make_sample(SAMPLE_B))
    (tmp_path / "c.docx").write_bytes(_make_sample(SAMPLE_C))

    result = TemplateBuilder().from_samples(
        tmp_path, template_id="derived", title="Derived"
    )
    assert result.template.template_id == "derived"
    titles = {s.title.lower() for s in result.template.all_sections()}
    assert any("introduction" in t for t in titles)
