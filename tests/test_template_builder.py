"""Tests for the Word .docx → JSON template builder."""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from services.template_service import DocxAdapter, DocxAdapterOptions, TemplateBuilder
from shared.schemas import (
    DataBindingType,
    GenerationMode,
    ReportTemplate,
    TemplateStatus,
)


def _build_sample_docx() -> bytes:
    """A tiny Word doc with realistic heading hierarchy for testing."""
    doc = Document()
    doc.add_heading("Investigator's Brochure — XYZ-001", level=0)  # Title style
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This Investigator's Brochure describes the investigational product XYZ-001."
    )
    doc.add_heading("2. Nonclinical Studies", level=1)
    doc.add_heading("2.1 Pharmacology", level=2)
    doc.add_paragraph(
        "Primary pharmacology was assessed in vitro and in murine xenograft models."
    )
    doc.add_heading("2.2 Toxicology", level=2)
    doc.add_paragraph(
        "Pivotal repeat-dose toxicology was performed in rat and dog species."
    )
    doc.add_heading("3. Effects in Humans", level=1)
    doc.add_heading("3.1 Pharmacokinetics in Humans", level=2)
    doc.add_heading("3.2 Safety and Efficacy", level=2)
    doc.add_paragraph(
        "Adverse events were graded per CTCAE; SAEs reported per ICH E2A."
    )
    doc.add_heading("4. Summary and Guidance for the Investigator", level=1)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_no_headings_docx() -> bytes:
    """A Word doc with NO heading styles — should produce zero sections."""
    doc = Document()
    doc.add_paragraph("Just a paragraph with no heading style applied.")
    doc.add_paragraph("Another paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_deep_nesting_docx() -> bytes:
    """Heading levels 1, 2, 3 to verify counter behaviour."""
    doc = Document()
    doc.add_heading("A", level=1)
    doc.add_heading("A1", level=2)
    doc.add_heading("A1a", level=3)
    doc.add_heading("A1b", level=3)
    doc.add_heading("A2", level=2)
    doc.add_heading("B", level=1)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- DocxAdapter -----------------------------------------------------------


def test_adapter_produces_draft_template() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="test_ib"))
    template = adapter.from_bytes(_build_sample_docx(), filename="sample.docx")
    assert isinstance(template, ReportTemplate)
    assert template.template_id == "test_ib"
    assert template.status == TemplateStatus.DRAFT
    assert template.metadata.source_origin == "from_docx"


def test_adapter_preserves_heading_hierarchy() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="test_ib"))
    template = adapter.from_bytes(_build_sample_docx())
    top_ids = [s.section_id for s in template.sections]
    # Title doesn't count; we expect "1", "2", "3", "4" at top level
    assert top_ids == ["1", "2", "3", "4"]
    sec2 = next(s for s in template.sections if s.section_id == "2")
    assert [c.section_id for c in sec2.children] == ["2.1", "2.2"]


def test_adapter_strips_explicit_heading_numbers_from_titles() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="t"))
    template = adapter.from_bytes(_build_sample_docx())
    sec1 = next(s for s in template.sections if s.section_id == "1")
    # Title was "1. Introduction" — number stripped
    assert sec1.title == "Introduction"
    sec22 = next(c for s in template.sections for c in s.children if c.section_id == "2.2")
    assert sec22.title == "Toxicology"


def test_adapter_handles_three_levels_of_nesting() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="t"))
    template = adapter.from_bytes(_build_deep_nesting_docx())
    top_ids = [s.section_id for s in template.sections]
    assert top_ids == ["1", "2"]  # "A" and "B"
    a = template.sections[0]
    assert [c.section_id for c in a.children] == ["1.1", "1.2"]
    a1 = a.children[0]
    assert [g.section_id for g in a1.children] == ["1.1.1", "1.1.2"]


def test_adapter_uses_content_hints_in_prompt() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="t"))
    template = adapter.from_bytes(_build_sample_docx())
    intro = next(s for s in template.sections if s.section_id == "1")
    assert intro.generation.mode == GenerationMode.LLM
    assert "content hint" in (intro.generation.prompt_template or "").lower()
    # The hint text must appear (or a prefix of it)
    assert "investigational product" in (intro.generation.prompt_template or "").lower()


def test_adapter_infers_file_set_tag_from_section_title() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="t"))
    template = adapter.from_bytes(_build_sample_docx())
    # 2.1 "Pharmacology" should pull a "pharmacology" tag
    pharm = next(c for s in template.sections for c in s.children if c.section_id == "2.1")
    tags = {
        tag
        for b in pharm.data_bindings
        if b.type == DataBindingType.FILE_SET
        for tag in getattr(b, "filter_tags", [])
    }
    assert "pharmacology" in tags

    # 2.2 "Toxicology" should pull a "toxicology" tag
    tox = next(c for s in template.sections for c in s.children if c.section_id == "2.2")
    tags = {
        tag
        for b in tox.data_bindings
        if b.type == DataBindingType.FILE_SET
        for tag in getattr(b, "filter_tags", [])
    }
    assert "toxicology" in tags


def test_adapter_always_includes_product_name_binding() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="t"))
    template = adapter.from_bytes(_build_sample_docx())
    for section in template.all_sections():
        ids = {b.binding_id for b in section.data_bindings}
        assert "product_name" in ids


def test_adapter_handles_no_headings_gracefully() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="t"))
    template = adapter.from_bytes(_build_no_headings_docx())
    assert template.sections == []


def test_adapter_validation_rules_default_to_strict() -> None:
    adapter = DocxAdapter(DocxAdapterOptions(template_id="t"))
    template = adapter.from_bytes(_build_sample_docx())
    intro = next(s for s in template.sections if s.section_id == "1")
    rule_codes = {r.rule for r in intro.validation_rules}
    assert "must_cite_every_number" in rule_codes
    assert "no_unbound_claims" in rule_codes


# --- TemplateBuilder façade -----------------------------------------------


def test_builder_returns_template_and_warnings(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    path.write_bytes(_build_sample_docx())
    result = TemplateBuilder().from_docx(path, template_id="test_ib")
    assert result.template.template_id == "test_ib"
    assert isinstance(result.warnings, tuple)


def test_builder_warns_on_no_headings(tmp_path: Path) -> None:
    path = tmp_path / "no_headings.docx"
    path.write_bytes(_build_no_headings_docx())
    result = TemplateBuilder().from_docx(path, template_id="empty_ib")
    assert any("verify the source" in w for w in result.warnings)


def test_builder_roundtrip_through_json(tmp_path: Path) -> None:
    """The drafted template must serialize to JSON and reload as the
    canonical ReportTemplate schema (no schema drift)."""
    path = tmp_path / "sample.docx"
    path.write_bytes(_build_sample_docx())
    result = TemplateBuilder().from_docx(path, template_id="rt_test")
    payload = result.template.model_dump_json()
    reloaded = ReportTemplate.model_validate_json(payload)
    assert reloaded.template_id == "rt_test"
    assert reloaded.all_sections() == result.template.all_sections()
