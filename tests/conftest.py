"""Shared pytest fixtures for the test suite."""

from __future__ import annotations

import io
import sys
from pathlib import Path

# Make the repo root importable when running pytest from any directory.
REPO_ROOT = Path(__file__).resolve().parent.parent
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))


def make_sample_docx() -> bytes:
    """Create a small in-memory DOCX with headings, paragraphs, and a table.

    Used by parser tests. Structure:
      H1: "Nonclinical Studies"
        H2: "Pharmacokinetics"
          paragraph + table
        H2: "Toxicology"
          paragraph
    """
    from docx import Document

    document = Document()
    document.add_heading("Nonclinical Studies", level=1)

    document.add_heading("Pharmacokinetics", level=2)
    document.add_paragraph(
        "The compound demonstrated linear pharmacokinetics across the dose range studied."
    )
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Cmax"
    table.rows[1].cells[1].text = "120 ng/mL"

    document.add_heading("Toxicology", level=2)
    document.add_paragraph("NOAEL was identified as 30 mg/kg/day in the 13-week rat study.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def make_sample_xlsx() -> bytes:
    """Create a small in-memory XLSX with one populated sheet."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "AE_Summary"
    ws["A1"] = "SOC"
    ws["B1"] = "Active n"
    ws["C1"] = "Placebo n"
    ws["A2"] = "Gastrointestinal disorders"
    ws["B2"] = 23
    ws["C2"] = 12
    ws["A3"] = "Nervous system disorders"
    ws["B3"] = 15
    ws["C3"] = 9

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_sample_pdf() -> bytes:
    """Create a small in-memory PDF using reportlab.

    Two pages: a title page and a body page with multiple paragraphs.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)

    # Page 1
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 720, "Investigator's Brochure — Compound XYZ-001")
    c.setFont("Helvetica", 11)
    c.drawString(72, 690, "Edition 1.0")
    c.showPage()

    # Page 2
    c.setFont("Helvetica", 11)
    c.drawString(72, 720, "The compound is a small-molecule inhibitor of kinase X.")
    c.drawString(72, 696, "")  # blank line separator
    c.drawString(72, 672, "It demonstrated activity in murine xenograft models at 10 mg/kg.")
    c.showPage()

    c.save()
    return buf.getvalue()
